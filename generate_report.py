from docx import Document
from db_connect import collection

records = list(collection.find({}, {"_id": 0}).sort("table_number", 1))

doc = Document()
doc.add_heading("Restaurant Table Management Report", level=1)
doc.add_paragraph(
    "This report was generated automatically from MongoDB using Python."
)

table = doc.add_table(rows=1, cols=6)
table.style = "Table Grid"

headers = [
    "Table No",
    "Capacity",
    "Location",
    "Status",
    "Customer Name",
    "Reservation Time"
]

for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header

for item in records:
    row = table.add_row().cells
    row[0].text = str(item.get("table_number", ""))
    row[1].text = str(item.get("capacity", ""))
    row[2].text = str(item.get("location", ""))
    row[3].text = str(item.get("status", ""))
    row[4].text = str(item.get("customer_name", ""))
    row[5].text = str(item.get("reservation_time", ""))

doc.save("restaurant_report.docx")
print("DOCX report created successfully.")
